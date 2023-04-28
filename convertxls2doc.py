import openpyxl
from docx import Document
from docx.shared import Inches

# Open the Excel file
workbook = openpyxl.load_workbook('ps.xlsx')

# Get the first sheet in the workbook
sheet = workbook.active

# Create a new Word document
document = Document()

# Loop through each row in the Excel sheet
for row in sheet.iter_rows():
    # Create a new paragraph in the Word document
    paragraph = document.add_paragraph()
    
    # Loop through each cell in the row
    for cell in row:
        # Add the cell value to the paragraph
        paragraph.add_run(str(cell.value))
        paragraph.add_run('\t') # add tab to separate the cells
    
    # Add a new line at the end of the paragraph
    paragraph.add_run('\n')

# Save the Word document
document.save('ps2.docx')
